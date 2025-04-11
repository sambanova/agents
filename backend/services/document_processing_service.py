import os
import tempfile
import base64
import mimetypes
from typing import List, Dict, Any
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument


class DocumentProcessingService:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def process_document(
        self, file_content: bytes, filename: str
    ) -> List[LangchainDocument]:
        """Process uploaded document and return chunks of text with metadata."""
        # Get file extension
        _, ext = os.path.splitext(filename.lower())

        # Create temp file to save content
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
            temp_file.write(file_content)
            temp_path = temp_file.name

        try:
            if ext in [".png", ".jpg", ".jpeg"]:
                base_64_url = self._base_64_data_url_from_doc(temp_path)   
                docs = [LangchainDocument(page_content="attached_image",
                                          metadata={"source": filename, "base_64_url": base_64_url})]
            else:
            # Extract text based on file type
                if ext == ".pdf":
                    text = self._extract_pdf_text(temp_path)
                elif ext in [".doc", ".docx"]:
                    text = self._extract_docx_text(temp_path)
                elif ext in [".csv", ".xlsx", ".xls"]:
                    text = self._extract_spreadsheet_text(temp_path)
                else:
                    raise ValueError(f"Unsupported file type: {ext}")
                # Split text into chunks
                docs = self.text_splitter.create_documents(
                    texts=[text], metadatas=[{"source": filename}]
                )

            return docs

        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.unlink(temp_path)

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_spreadsheet_text(self, file_path: str) -> str:
        """Extract text from spreadsheet files."""
        if file_path.endswith(".csv"):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        # Convert dataframe to string representation
        return df.to_string()
        
    def _base_64_data_url_from_doc(self, file_path: str, add_mimetype: bool=True) -> str:
        """Convert to to Base64 URL data"""
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type is None:
            
            raise ValueError("Could not determine MIME type of the file")

        with open(file_path, "rb") as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode("utf-8")

        if add_mimetype:
            data_url = f"data:{mime_type};base64,{encoded_string}"
        else:
            data_url = encoded_string
        return data_url